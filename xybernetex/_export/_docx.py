"""
DOCX renderer — converts an Artifact to a Word document.
Requires: pip install xybernetex[docx]   (python-docx)
"""
from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from xybernetex._models import Artifact

_MISSING = (
    "python-docx is required for DOCX export.\n"
    "Install it with:  pip install xybernetex[docx]"
)


def render_docx(artifact: "Artifact", path: str) -> str:
    """
    Write *artifact* to a Word document at *path*.

    If *path* is a directory the file is named after the artifact title.
    Returns the path of the written file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError(_MISSING) from None

    from xybernetex._export._markdown import parse, inline_runs
    from xybernetex._models import _safe_filename

    dest = _resolve_path(path, artifact.title, ".docx")

    doc = Document()

    # ── Document title ─────────────────────────────────────────────────────────
    title_para = doc.add_heading(artifact.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Subtle sub-line with artifact type
    sub = doc.add_paragraph(artifact.artifact_type.replace("_", " ").title())
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    sub.runs[0].font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # ── Body ───────────────────────────────────────────────────────────────────
    for block in parse(artifact.content):
        btype = block["type"]

        if btype == "heading":
            doc.add_heading(block["text"], level=min(block["level"], 4))

        elif btype == "paragraph":
            para = doc.add_paragraph()
            for run in inline_runs(block["text"]):
                r = para.add_run(run["text"])
                r.bold = run["bold"]
                r.italic = run["italic"]
                if run["code"]:
                    r.font.name = "Courier New"
                    r.font.size = Pt(9)

        elif btype == "code":
            para = doc.add_paragraph(block["text"])
            para.style = "No Spacing"
            for r in para.runs:
                r.font.name = "Courier New"
                r.font.size = Pt(9)
            # Light grey shading
            _shade_paragraph(para, "F2F2F2")

        elif btype == "bullet_list":
            for item in block["items"]:
                para = doc.add_paragraph(style="List Bullet")
                for run in inline_runs(item):
                    r = para.add_run(run["text"])
                    r.bold = run["bold"]
                    r.italic = run["italic"]

        elif btype == "ordered_list":
            for item in block["items"]:
                para = doc.add_paragraph(style="List Number")
                for run in inline_runs(item):
                    r = para.add_run(run["text"])
                    r.bold = run["bold"]
                    r.italic = run["italic"]

        elif btype == "table":
            headers = block["headers"]
            rows = block["rows"]
            if not headers:
                continue
            col_count = max(len(headers), max((len(r) for r in rows), default=0))
            table = doc.add_table(rows=1 + len(rows), cols=col_count)
            table.style = "Light List Accent 1"

            # Header row
            hdr_cells = table.rows[0].cells
            for ci, h in enumerate(headers[:col_count]):
                hdr_cells[ci].text = h
                for run in hdr_cells[ci].paragraphs[0].runs:
                    run.bold = True

            # Data rows
            for ri, row_data in enumerate(rows):
                cells = table.rows[ri + 1].cells
                for ci, val in enumerate(row_data[:col_count]):
                    cells[ci].text = val

        elif btype == "hr":
            doc.add_paragraph("─" * 60)

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    return str(dest)


# ── Helpers ────────────────────────────────────────────────────────────────────

def _resolve_path(path: str, title: str, ext: str) -> Path:
    from xybernetex._models import _safe_filename
    p = Path(path)
    if p.is_dir():
        p = p / f"{_safe_filename(title)}{ext}"
    return p


def _shade_paragraph(para, hex_color: str) -> None:
    """Apply a background shade to a paragraph via raw XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)
